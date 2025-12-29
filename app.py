import streamlit as st
import openai
import httpx
import requests
from bs4 import BeautifulSoup
from docx import Document
import PyPDF2
import pandas as pd
import os
import pythoncom
from doc2docx import convert
import re

# --- НАСТРОЙКИ ДОСТУПА ---
API_KEY = "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"
# Новый SOCKS5 прокси
PROXY_URL = "socks5://YtvW3X:MgRYbP@45.91.209.157:12782"

# Настройка клиента с поддержкой SOCKS5 и долгими таймаутами
custom_http_client = httpx.Client(
    proxy=PROXY_URL,
    timeout=httpx.Timeout(600.0, connect=60.0, read=540.0),
    trust_env=False # Полный игнор системных прокси
)

ai_client = openai.OpenAI(
    api_key=API_KEY,
    http_client=custom_http_client
)

st.set_page_config(page_title="AI-Тендерный отдел v2.1", layout="wide")
st.title("🤖 AI-Тендерный отдел v2.1")

if not os.path.exists("temp"):
    os.makedirs("temp")

# --- ФУНКЦИИ ОБРАБОТКИ ТЕКСТА ---
def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text(file_path, file_name):
    ext = file_name.split('.')[-1].lower()
    text = ""
    try:
        if ext == 'docx':
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == 'doc':
            pythoncom.CoInitialize()
            docx_path = file_path + "x"
            convert(file_path, doc_path)
            doc = Document(docx_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            if os.path.exists(docx_path): os.remove(docx_path)
        elif ext == 'pdf':
            reader = PyPDF2.PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == 'xlsx':
            df = pd.read_excel(file_path)
            text = f"Таблица {file_name}:\n" + df.to_string()
    except Exception as e:
        st.error(f"Ошибка чтения {file_name}: {e}")
    return clean_text(text)

# --- МОДУЛЬ ЕИС (БЕЗ ПРОКСИ) ---
def download_eis_files(url):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120.0.0.0'}
    files_data = []
    session = requests.Session()
    session.trust_env = False 
    try:
        if "common-info" in url:
            url = url.replace("common-info", "documents")
        response = session.get(url, headers=headers, timeout=20)
        soup = BeautifulSoup(response.text, 'html.parser')
        links = soup.find_all('a', href=True)
        download_links = [l['href'] for l in links if "download" in l['href'].lower()]
        
        if not download_links:
            st.warning("Файлы не найдены.")
            return []

        progress_bar = st.progress(0)
        to_download = download_links[:5] 
        for i, link in enumerate(to_download):
            full_link = link if link.startswith('http') else "https://zakupki.gov.ru" + link
            f_resp = session.get(full_link, headers=headers)
            d_header = f_resp.headers.get('content-disposition', '')
            fname = re.findall("filename=(.+)", d_header)
            fname = fname[0].strip('"') if fname else f"doc_{i}.docx"
            fname = "".join([c for c in fname if c.isalnum() or c in "._- "]).strip()
            f_path = os.path.join("temp", fname)
            with open(f_path, "wb") as f:
                f.write(f_resp.content)
            files_data.append({"path": f_path, "name": fname})
            progress_bar.progress((i + 1) / len(to_download))
        return files_data
    except Exception as e:
        st.error(f"Ошибка ЕИС: {e}")
        return []

# --- ГЛУБОКИЙ АНАЛИЗ (КАРТА РИСКОВ) ---
def run_ai_analysis(context_text):
    if not context_text.strip():
        st.error("Текст документации пуст.")
        return

    with st.spinner("ИИ заменяет тендерный отдел: глубокий аудит..."):
        try:
            safe_text = context_text[:120000]
            prompt = f"""
            Действуй как ИИ-ассистент тендерного отдела. Твоя цель — снять операционную нагрузку с руководителя.
            Проанализируй документацию по следующим критериям:

            1. **Умный фильтр (неформализованные данные)**: 
               - Количество поставок, формат упаковки, полный список необходимых документов.
            
            2. **Карта рисков и закрывашек**:
               - Реалистичность характеристик: есть ли бренды или "заточки" под конкретного производителя?.
               - Сверка ТЗ с типичными условиями: избыточные требования или скрытые штрафы.
            
            3. **Тендерная корзина (Расчет)**:
               - Составь четкий список товаров для поиска у поставщиков.
            
            4. **Контроль и Сроки**:
               - Проверь сроки поставки и оплаты. Насколько они объективны?.

            ТЕКСТ ДЛЯ АНАЛИЗА:
            {safe_text}
            """

            response = ai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": "Ты высококвалифицированный тендерный аналитик в РФ."},
                          {"role": "user", "content": prompt}]
            )
            st.markdown("### 📊 Аналитическое резюме")
            st.markdown(response.choices[0].message.content)
            
        except Exception as e:
            st.error(f"Ошибка связи (SOCKS5): {e}")
            st.info("Проверьте, установлен ли пакет httpx[socks] через pip.")

# --- ИНТЕРФЕЙС ---
t1, t2 = st.tabs(["📁 Ручная загрузка", "🔗 Ссылка ЕИС"])

with t1:
    up = st.file_uploader("Документы (PDF, Word, Excel)", accept_multiple_files=True)
    if up and st.button("🚀 Начать анализ"):
        all_c = ""
        pb = st.progress(0)
        for i, f in enumerate(up):
            p = os.path.join("temp", f.name)
            with open(p, "wb") as t: t.write(f.getbuffer())
            all_c += f"\n\n=== {f.name} ===\n" + extract_text(p, f.name)
            os.remove(p)
            pb.progress((i+1)/len(up))
        run_ai_analysis(all_c)

with t2:
    url = st.text_input("Вставьте ссылку на тендер из ЕИС:")
    if st.button("🔍 Загрузить и проверить"):
        files = download_eis_files(url)
        if files:
            full_t = ""
            for fi in files:
                full_t += f"\n\n=== {fi['name']} ===\n" + extract_text(fi['path'], fi['name'])
                os.remove(fi['path'])
            run_ai_analysis(full_t)